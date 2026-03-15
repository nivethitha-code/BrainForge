from docx import Document
import os

def convert_docx_to_md(docx_path, output_path):
    doc = Document(docx_path)
    with open(output_path, 'w', encoding='utf-8') as f:
        for para in doc.paragraphs:
            if para.style.name.startswith('Heading'):
                level = para.style.name.split()[-1]
                if level.isdigit():
                    f.write('#' * int(level) + ' ' + para.text + '\n\n')
                else:
                    f.write('# ' + para.text + '\n\n')
            else:
                f.write(para.text + '\n\n')
        
        for i, table in enumerate(doc.tables):
            f.write(f'### Table {i+1}\n\n')
            for row in table.rows:
                row_text = ' | '.join(cell.text.strip().replace('\n', ' ') for cell in row.cells)
                f.write('| ' + row_text + ' |\n')
            f.write('\n')

if __name__ == "__main__":
    convert_docx_to_md(r'd:\Semester-6\PROJECTS\Quizz App\BrainForge_implementation_plan.docx', r'd:\Semester-6\PROJECTS\Quizz App\plan.md')
